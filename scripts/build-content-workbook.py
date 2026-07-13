#!/usr/bin/env python3
"""Build an editable Word workbook from the Sanctuary Explorer content modules."""

from __future__ import annotations

import json
import os
import re
import subprocess
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import html


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "editorial" / "Sanctuary_Explorer_Content_Workbook.docx"
FIELD_MAP_OUTPUT = ROOT / "editorial" / "Sanctuary_Explorer_Content_Field_Map.json"
NODE = Path(
    os.environ.get(
        "SANCTUARY_NODE",
        "/Users/samuel/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/bin/node",
    )
)

PRESET = {
    "font": "Calibri",
    "body_size": 11,
    "body_after": 6,
    "body_line": 1.25,
    "heading_1": {"size": 16, "color": "2E74B5", "before": 18, "after": 10},
    "heading_2": {"size": 13, "color": "2E74B5", "before": 14, "after": 7},
    "heading_3": {"size": 12, "color": "1F4D78", "before": 10, "after": 5},
}

CONTENT_SPECS = [
    {
        "part": "Live Website Copy",
        "title": "Site Shell and Navigation",
        "module": "site-copy",
        "source": "content/site-copy.js",
        "exports": ["siteCopy"],
        "note": "Document metadata, navigation labels, headings, help text, accessibility labels, and footer copy.",
    },
    {
        "part": "Live Website Copy",
        "title": "Sanctuary Articles and Restoration Path",
        "module": "articles",
        "source": "content/articles.js",
        "exports": ["articles", "articleEnhancements", "phases"],
        "note": "Furniture, structure, process, fulfillment, application, summary, key-text, and restoration-path copy.",
    },
    {
        "part": "Live Website Copy",
        "title": "Sanctuary Ministry",
        "module": "ministry",
        "source": "content/ministry.js",
        "exports": ["ministryViews", "offerings", "specialCeremonies", "serviceStudies"],
        "note": "Sacrifices, special ceremonies, daily service, yearly service, and ministry tab copy.",
    },
    {
        "part": "Live Website Copy",
        "title": "Priestly Attire",
        "module": "attire",
        "source": "content/attire.js",
        "exports": ["attireGroups", "priestlyAttire"],
        "note": "Garment names, descriptions, meanings, typology, Scripture references, and applications.",
    },
    {
        "part": "Live Website Copy",
        "title": "Festival Calendar",
        "module": "feasts",
        "source": "content/feasts.js",
        "exports": ["feasts"],
        "note": "Feast names, dates, movements, procedures, significance, fulfillment, and study notes.",
    },
    {
        "part": "Live Website Copy",
        "title": "Sanctuary Chronicles",
        "module": "chronicles",
        "source": "content/chronicles.js",
        "exports": [
            "chroniclePerspectiveInfo",
            "documentChronicleStories",
            "documentChronicleCompanionNotes",
        ],
        "note": "The active priest/commoner stories, scene narratives, Scripture captions, and companion notes.",
    },
    {
        "part": "Live Website Copy",
        "title": "1844 Focus Study",
        "module": "focus-1844",
        "source": "content/focus-1844.js",
        "exports": [
            "focusTimeline",
            "focusStudyCards",
            "focusHeroChain",
            "focusLogicSteps",
            "focusDanielLinkSteps",
            "focusChartSegments",
            "focusAssuranceCards",
        ],
        "note": "Timeline, study cards, prophecy chain, logic flow, chart labels, and assurance copy.",
    },
    {
        "part": "Live Website Copy",
        "title": "FAQ Study Hub",
        "module": "faq",
        "source": "content/faq.js",
        "exports": ["polishedFaqArticles"],
        "note": "The active FAQ articles, quick answers, study sections, key texts, and related-link labels.",
    },
    {
        "part": "Live Website Copy",
        "title": "3D Guided Paths and Cinematic Scenes",
        "module": "scene-paths",
        "source": "content/scene-paths.js",
        "exports": ["STUDY_PATHS", "CINEMATIC_PATHS"],
        "note": "Guided-path captions and cinematic titles, narration, and Scripture references. Camera and actor data are intentionally omitted.",
    },
    {
        "part": "Legacy Source Appendix",
        "title": "Chronicle Source Drafts",
        "module": "chronicles",
        "source": "content/chronicles.js",
        "exports": ["chronicleStories", "chronicleCompanionNotes"],
        "note": "Earlier Chronicle seed copy retained in the codebase but replaced at runtime by the active document Chronicle exports.",
    },
    {
        "part": "Legacy Source Appendix",
        "title": "FAQ Source Drafts",
        "module": "faq",
        "source": "content/faq.js",
        "exports": ["faqArticles"],
        "note": "Earlier FAQ source copy retained in the codebase. The live FAQ uses polishedFaqArticles above.",
    },
]

EXPORT_TITLES = {
    "siteCopy": "Site Copy",
    "articles": "Sanctuary Articles",
    "articleEnhancements": "Article Enhancements",
    "phases": "Path to Restoration",
    "ministryViews": "Ministry Tab Labels",
    "offerings": "Sacrifices and Offerings",
    "specialCeremonies": "Special Ceremonies",
    "serviceStudies": "Daily and Yearly Services",
    "attireGroups": "Attire Group Labels",
    "priestlyAttire": "Garment Explanations",
    "feasts": "Feasts",
    "chroniclePerspectiveInfo": "Perspective Selector",
    "documentChronicleStories": "Active Chronicle Stories",
    "documentChronicleCompanionNotes": "Active Chronicle Companion Notes",
    "focusTimeline": "Timeline",
    "focusStudyCards": "Study Cards",
    "focusHeroChain": "Hero Chain",
    "focusLogicSteps": "Sanctuary Logic Flow",
    "focusDanielLinkSteps": "Daniel and Revelation Link",
    "focusChartSegments": "Prophecy Chart Labels",
    "focusAssuranceCards": "Assurance Cards",
    "polishedFaqArticles": "Active FAQ Articles",
    "STUDY_PATHS": "Guided Study Paths",
    "CINEMATIC_PATHS": "Cinematic Scene Captions",
    "chronicleStories": "Legacy Chronicle Stories",
    "chronicleCompanionNotes": "Legacy Chronicle Companion Notes",
    "faqArticles": "Legacy FAQ Articles",
}

SKIP_KEYS = {
    "id",
    "articleId",
    "icon",
    "iconKey",
    "color",
    "accent",
    "layout",
    "fill",
    "left",
    "width",
    "camera",
    "target",
    "focusId",
    "actors",
    "actorStates",
    "transition",
    "hold",
    "duration",
    "position",
    "rotation",
    "scale",
    "action",
    "defaultGarment",
    "group",
    "zone",
    "perspective",
    "serviceType",
    "section",
    "href",
    "url",
    "featured",
    "selector",
}


def load_content() -> list[dict[str, Any]]:
    node_specs = [
        {"module": spec["module"], "exports": spec["exports"]}
        for spec in CONTENT_SPECS
    ]
    script = f"""
const specs = {json.dumps(node_specs)};
const cache = new Map();
const result = [];
for (const spec of specs) {{
  let moduleValue = cache.get(spec.module);
  if (!moduleValue) {{
    moduleValue = await import(`./content/${{spec.module}}.js`);
    cache.set(spec.module, moduleValue);
  }}
  const exportsValue = {{}};
  for (const exportName of spec.exports) exportsValue[exportName] = moduleValue[exportName];
  result.push(exportsValue);
}}
process.stdout.write(JSON.stringify(result));
"""
    completed = subprocess.run(
        [str(NODE), "--input-type=module", "-e", script],
        cwd=ROOT,
        check=True,
        capture_output=True,
        text=True,
    )
    return json.loads(completed.stdout)


def set_cell_or_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = PRESET["font"]
    normal._element.rPr.rFonts.set(qn("w:ascii"), PRESET["font"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["font"])
    normal.font.size = Pt(PRESET["body_size"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(PRESET["body_after"])
    normal.paragraph_format.line_spacing = PRESET["body_line"]

    for style_name, token_name in (
        ("Heading 1", "heading_1"),
        ("Heading 2", "heading_2"),
        ("Heading 3", "heading_3"),
    ):
        style = styles[style_name]
        token = PRESET[token_name]
        style.font.name = PRESET["font"]
        style._element.rPr.rFonts.set(qn("w:ascii"), PRESET["font"])
        style._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["font"])
        style.font.size = Pt(token["size"])
        style.font.color.rgb = RGBColor.from_string(token["color"])
        style.font.bold = True
        style.paragraph_format.space_before = Pt(token["before"])
        style.paragraph_format.space_after = Pt(token["after"])
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    group_style = styles.add_style("Content Group", WD_STYLE_TYPE.PARAGRAPH)
    group_style.base_style = normal
    group_style.font.name = PRESET["font"]
    group_style._element.rPr.rFonts.set(qn("w:ascii"), PRESET["font"])
    group_style._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["font"])
    group_style.font.size = Pt(11)
    group_style.font.bold = True
    group_style.font.color.rgb = RGBColor.from_string("2E74B5")
    group_style.paragraph_format.space_before = Pt(8)
    group_style.paragraph_format.space_after = Pt(3)
    group_style.paragraph_format.line_spacing = 1.0
    group_style.paragraph_format.keep_with_next = True

    field_style = styles.add_style("Content Field", WD_STYLE_TYPE.PARAGRAPH)
    field_style.base_style = normal
    field_style.font.name = PRESET["font"]
    field_style._element.rPr.rFonts.set(qn("w:ascii"), PRESET["font"])
    field_style._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["font"])
    field_style.font.size = Pt(11)
    field_style.paragraph_format.space_before = Pt(0)
    # Named compact-field override: the manuscript carries more than 2,800
    # individually mapped fields, so field paragraphs use a tighter rhythm.
    field_style.paragraph_format.space_after = Pt(5)
    field_style.paragraph_format.line_spacing = 1.15
    field_style.paragraph_format.widow_control = True

    note_style = styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
    note_style.base_style = normal
    note_style.font.name = PRESET["font"]
    note_style._element.rPr.rFonts.set(qn("w:ascii"), PRESET["font"])
    note_style._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["font"])
    note_style.font.size = Pt(8.5)
    note_style.font.color.rgb = RGBColor.from_string("64748B")
    note_style.paragraph_format.space_before = Pt(0)
    note_style.paragraph_format.space_after = Pt(6)
    note_style.paragraph_format.line_spacing = 1.0

    callout_style = styles.add_style("Editorial Callout", WD_STYLE_TYPE.PARAGRAPH)
    callout_style.base_style = normal
    callout_style.font.name = PRESET["font"]
    callout_style._element.rPr.rFonts.set(qn("w:ascii"), PRESET["font"])
    callout_style._element.rPr.rFonts.set(qn("w:hAnsi"), PRESET["font"])
    callout_style.font.size = Pt(10.5)
    callout_style.paragraph_format.left_indent = Inches(0.18)
    callout_style.paragraph_format.right_indent = Inches(0.18)
    callout_style.paragraph_format.space_before = Pt(6)
    callout_style.paragraph_format.space_after = Pt(10)
    callout_style.paragraph_format.line_spacing = 1.2


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = p.add_run("SANCTUARY EXPLORER")
    set_run_font(left, size=8, color="64748B", bold=True)
    right = p.add_run("\tCONTENT WORKBOOK")
    set_run_font(right, size=8, color="64748B")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Page ")
    set_run_font(run, size=8, color="64748B")
    field_run = p.add_run()
    set_run_font(field_run, size=8, color="64748B")
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    field_run._r.append(fld_char_begin)
    field_run._r.append(instr_text)
    field_run._r.append(fld_char_separate)
    field_run._r.append(text_node)
    field_run._r.append(fld_char_end)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(10)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("EDITORIAL MASTER")
    set_run_font(run, size=10, color="A86718", bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Sanctuary Explorer")
    set_run_font(run, size=30, color="17213A", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("Complete Website Content Workbook")
    set_run_font(run, size=15, color="2B5163")

    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_after = Pt(70)
    run = descriptor.add_run("Current wording, exact content keys, and a replacement prompt")
    set_run_font(run, size=10.5, color="64748B", italic=True)

    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_line.paragraph_format.space_after = Pt(3)
    run = date_line.add_run("Prepared July 13, 2026")
    set_run_font(run, size=10, color="17213A", bold=True)

    status = doc.add_paragraph()
    status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = status.add_run("First-pass extraction: wording unchanged")
    set_run_font(run, size=9.5, color="64748B")

    doc.add_page_break()


def add_front_matter(doc: Document, field_count: int, word_count: int) -> None:
    doc.add_heading("How to Use This Workbook", level=1)
    doc.add_paragraph(
        "This document collects the current website wording in one editable manuscript. "
        "The live copy appears first. Older source drafts that remain in the codebase appear in the final appendix."
    )

    p = doc.add_paragraph(style="Editorial Callout")
    set_cell_or_paragraph_shading(p, "F4F6F9")
    label = p.add_run("Editing rule\n")
    set_run_font(label, size=10.5, color="1F4D78", bold=True)
    body = p.add_run(
        "Change the wording after each short marker, such as [[FIELD-ID]]. Keep every marker exactly as written. "
        "The markers are the map I will use to place your revised text back into the correct article, panel, or scene."
    )
    set_run_font(body, size=10.5, color="17213A")

    doc.add_heading("What You May Edit", level=2)
    doc.add_paragraph(
        "You may rewrite titles, summaries, explanations, article paragraphs, Scripture captions, scene narration, "
        "button labels, help text, and accessibility wording. Section headings in this Word file are only navigation aids; "
        "the marked fields are the actual website copy."
    )

    doc.add_heading("What Has Been Protected", level=2)
    doc.add_paragraph(
        "Internal IDs, object mappings, CSS selectors, colors, icons, camera positions, actor instructions, timing values, "
        "and other presentation settings are omitted from the manuscript. Your rewrite therefore cannot accidentally alter the 3D scene or routing logic."
    )

    stats = doc.add_paragraph(style="Source Note")
    stats.add_run(f"Workbook scope: {field_count:,} editable text fields | approximately {word_count:,} words of current copy.")

    doc.add_heading("Prompt for Rewriting the Text", level=1)
    intro = doc.add_paragraph(
        "Attach this workbook to the writing tool you want to use, then paste the prompt below. Add any special editorial direction in the final line."
    )
    intro.paragraph_format.space_after = Pt(8)

    prompt_text = (
        "I am revising the text of the Sanctuary Explorer website. The attached Word workbook contains the current website copy. "
        "Rewrite it in language that is clear, devotional, academically responsible, beginner-friendly, and faithful to the sanctuary message from a Seventh-day Adventist perspective.\n\n"
        "Preserve every short field marker, such as [[FIELD-ID]], exactly. Rewrite only the wording that follows each marker. Do not rename, delete, reorder, or invent field markers. "
        "Keep short interface labels short. Keep 3D scene captions vivid and concise. Use longer explanation only where the workbook contains an article or study section. "
        "Preserve Scripture references and the KJV designation where it already appears unless a correction is clearly necessary. Do not introduce claims that cannot be supported from Scripture. "
        "Avoid repetition, religious cliches, unexplained jargon, and language that assumes the reader already understands the sanctuary. Maintain theological continuity across articles, scenes, FAQ answers, feasts, priestly attire, and the 1844 study.\n\n"
        "Return the complete revised text with all field markers intact, including fields you decide not to change. Do not summarize the revision and do not omit unchanged sections.\n\n"
        "Additional editorial direction: [Write my specific goals, audience, preferred tone, or doctrinal emphasis here.]"
    )
    p = doc.add_paragraph(style="Editorial Callout")
    set_cell_or_paragraph_shading(p, "E8EEF5")
    p.paragraph_format.line_spacing = 1.08
    set_run_font(p.add_run(prompt_text), size=9.5, color="17213A")


def humanize(value: str) -> str:
    value = re.sub(r"([a-z0-9])([A-Z])", r"\1 \2", value)
    value = value.replace("_", " ").replace("-", " ")
    return " ".join(value.split()).title()


def append_path(path: str, key: str | int) -> str:
    if isinstance(key, int):
        return f"{path}[{key}]"
    if re.fullmatch(r"[A-Za-z_$][A-Za-z0-9_$]*", key):
        return f"{path}.{key}"
    return f"{path}[{json.dumps(key, ensure_ascii=False)}]"


def is_skipped(key: str | None, value: Any, ancestors: Iterable[str]) -> bool:
    if key in SKIP_KEYS:
        return True
    if any(ancestor in SKIP_KEYS for ancestor in ancestors):
        return True
    if isinstance(value, str):
        stripped = value.strip()
        if stripped.startswith("<svg") or stripped.startswith("linear-gradient"):
            return True
        if re.fullmatch(r"#[0-9A-Fa-f]{3,8}", stripped):
            return True
    return False


def walk_strings(
    value: Any,
    path: str,
    *,
    key: str | None = None,
    ancestors: tuple[str, ...] = (),
) -> Iterable[tuple[str, str]]:
    if is_skipped(key, value, ancestors):
        return
    if isinstance(value, str):
        if value.strip():
            yield path, value
        return
    if isinstance(value, list):
        for index, item in enumerate(value):
            yield from walk_strings(
                item,
                append_path(path, index),
                key=str(index),
                ancestors=ancestors,
            )
        return
    if isinstance(value, dict):
        for child_key, child_value in value.items():
            if "selector" in value and child_key == "name":
                continue
            yield from walk_strings(
                child_value,
                append_path(path, child_key),
                key=child_key,
                ancestors=ancestors + ((key,) if key else ()),
            )


def item_label(item: Any, fallback: str) -> str:
    if isinstance(item, dict):
        if "selector" in item and isinstance(item.get("value"), str):
            return item["value"].strip()
        for key in ("title", "name", "label", "heading", "serviceLabel", "tab", "date", "reference", "id"):
            value = item.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()
    return fallback


def add_text_run(paragraph, text_value: str, *, bold: bool = False, italic: bool = False, color: str = "17213A") -> None:
    chunks = text_value.split("\n")
    for index, chunk in enumerate(chunks):
        if index:
            paragraph.add_run().add_break()
        if chunk:
            run = paragraph.add_run(chunk)
            set_run_font(run, size=11, color=color, bold=bold, italic=italic)


def append_html_content(paragraph, value: str) -> None:
    if not re.search(r"<[A-Za-z][^>]*>", value):
        add_text_run(paragraph, value)
        return

    root = html.fragment_fromstring(value, create_parent="div")

    def render_node(node, bold: bool = False, italic: bool = False) -> None:
        tag = node.tag.lower() if isinstance(node.tag, str) else ""
        child_bold = bold or tag in {"strong", "b"}
        child_italic = italic or tag in {"em", "i"}
        if tag == "br":
            paragraph.add_run().add_break()
        if node.text:
            add_text_run(paragraph, node.text, bold=child_bold, italic=child_italic)
        for child in node:
            render_node(child, child_bold, child_italic)
            if child.tail:
                add_text_run(paragraph, child.tail, bold=child_bold, italic=child_italic)

    if root.text:
        add_text_run(paragraph, root.text)
    for child in root:
        render_node(child)
        if child.tail:
            add_text_run(paragraph, child.tail)


def add_field(doc: Document, field_id: str, value: str) -> None:
    p = doc.add_paragraph(style="Content Field")
    marker = p.add_run(f"[[{field_id}]] ")
    set_run_font(marker, name="Consolas", size=8, color="64748B", bold=True)
    append_html_content(p, value)


def add_group_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph(title, style="Content Group")
    p.paragraph_format.keep_with_next = True


def render_nested(
    doc: Document,
    source: str,
    value: Any,
    path: str,
    *,
    field_ids: dict[tuple[str, str], str],
    key: str | None = None,
    ancestors: tuple[str, ...] = (),
    depth: int = 0,
) -> None:
    if is_skipped(key, value, ancestors):
        return
    if isinstance(value, str):
        if value.strip():
            add_field(doc, field_ids[(source, path)], value)
        return
    if isinstance(value, list):
        if not value:
            return
        for index, item in enumerate(value):
            item_path = append_path(path, index)
            is_site_copy_entry = isinstance(item, dict) and "selector" in item and "value" in item
            if isinstance(item, (dict, list)) and not is_site_copy_entry:
                label = item_label(item, f"{humanize(key or 'Item')} {index + 1}")
                add_group_heading(doc, f"{humanize(key or 'Item')} {index + 1}: {label}")
            render_nested(
                doc,
                source,
                item,
                item_path,
                field_ids=field_ids,
                key=str(index),
                ancestors=ancestors,
                depth=depth + 1,
            )
        return
    if isinstance(value, dict):
        for child_key, child_value in value.items():
            if "selector" in value and child_key == "name":
                continue
            if is_skipped(child_key, child_value, ancestors + ((key,) if key else ())):
                continue
            child_path = append_path(path, child_key)
            if isinstance(child_value, (dict, list)) and child_value:
                add_group_heading(doc, humanize(child_key))
            render_nested(
                doc,
                source,
                child_value,
                child_path,
                field_ids=field_ids,
                key=child_key,
                ancestors=ancestors + ((key,) if key else ()),
                depth=depth + 1,
            )


def render_export(
    doc: Document,
    source: str,
    export_name: str,
    value: Any,
    field_ids: dict[tuple[str, str], str],
) -> None:
    doc.add_heading(EXPORT_TITLES.get(export_name, humanize(export_name)), level=2)
    source_note = doc.add_paragraph(style="Source Note")
    source_note.add_run(f"Source export: {source}::{export_name}")

    if isinstance(value, dict):
        for item_key, item in value.items():
            if is_skipped(item_key, item, (export_name,)):
                continue
            label = item_label(item, humanize(item_key))
            doc.add_heading(f"{label}  [{item_key}]", level=3)
            render_nested(
                doc,
                source,
                item,
                append_path(export_name, item_key),
                field_ids=field_ids,
                key=item_key,
                ancestors=(export_name,),
            )
        return

    if isinstance(value, list):
        for index, item in enumerate(value):
            label = item_label(item, f"Item {index + 1}")
            item_id = item.get("id") if isinstance(item, dict) else None
            suffix = f"  [{item_id}]" if isinstance(item_id, str) else f"  [item {index + 1}]"
            doc.add_heading(f"{label}{suffix}", level=3)
            render_nested(
                doc,
                source,
                item,
                append_path(export_name, index),
                field_ids=field_ids,
                key=str(index),
                ancestors=(export_name,),
            )
        return

    render_nested(doc, source, value, export_name, field_ids=field_ids, key=export_name)


def build_workbook(content_values: list[dict[str, Any]]) -> tuple[int, int]:
    field_records: list[dict[str, str]] = []
    for spec, exports_value in zip(CONTENT_SPECS, content_values):
        for export_name, value in exports_value.items():
            for path, field_value in walk_strings(value, export_name, key=export_name):
                field_records.append(
                    {
                        "part": spec["part"],
                        "section": spec["title"],
                        "source": spec["source"],
                        "path": path,
                        "currentText": field_value,
                    }
                )
    for index, record in enumerate(field_records, start=1):
        record["fieldId"] = f"F{index:04d}"
    field_ids = {(record["source"], record["path"]): record["fieldId"] for record in field_records}
    field_count = len(field_records)
    word_count = sum(
        len(re.findall(r"\b[\w'-]+\b", html.fromstring(f"<div>{record['currentText']}</div>").text_content()))
        for record in field_records
    )

    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = "Sanctuary Explorer Complete Website Content Workbook"
    doc.core_properties.subject = "Editable website copy mapped to Sanctuary Explorer content modules"
    doc.core_properties.author = "Sanctuary Explorer"
    doc.core_properties.keywords = "Sanctuary Explorer, editorial workbook, website copy"
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings._element.append(update_fields)

    add_cover(doc)
    add_front_matter(doc, field_count, word_count)

    current_part = None
    for index, (spec, exports_value) in enumerate(zip(CONTENT_SPECS, content_values)):
        if spec["part"] != current_part:
            doc.add_page_break()
            current_part = spec["part"]
            doc.add_heading(current_part, level=1)
            if current_part == "Legacy Source Appendix":
                p = doc.add_paragraph(style="Editorial Callout")
                set_cell_or_paragraph_shading(p, "FFF8E8")
                set_run_font(
                    p.add_run(
                        "The following copy remains in the repository as source material but is not the active text shown on the live site. "
                        "Rewrite it only if you also want these retained drafts updated."
                    ),
                    size=10.5,
                    color="7A5A00",
                )

        doc.add_heading(spec["title"], level=1)
        note = doc.add_paragraph(style="Source Note")
        note.add_run(f"{spec['source']} | {spec['note']}")
        for export_name, value in exports_value.items():
            render_export(doc, spec["source"], export_name, value, field_ids)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    FIELD_MAP_OUTPUT.write_text(
        json.dumps(field_records, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return field_count, word_count


def main() -> None:
    content_values = load_content()
    field_count, word_count = build_workbook(content_values)
    print(f"Wrote {OUTPUT}")
    print(f"Wrote {FIELD_MAP_OUTPUT}")
    print(f"Editable fields: {field_count}")
    print(f"Approximate words: {word_count}")


if __name__ == "__main__":
    main()
